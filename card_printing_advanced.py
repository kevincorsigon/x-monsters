"""
Script para criar folhas de impressÃ£o personalizadas das cartas X-MONSTERS
Permite ajustar dimensÃµes e layout conforme necessÃ¡rio
"""

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
from math import floor

class CardPrintingGenerator:
    def __init__(self):
        # ConfiguraÃ§Ãµes padrÃ£o (podem ser modificadas)
        self.card_width_mm = 63
        self.card_height_mm = 88
        self.a4_width_mm = 210
        self.a4_height_mm = 297
        self.margin_mm = 10
        self.spacing_mm = 8  # EspaÃ§o entre cartas para facilitar recorte
        
    def calculate_layout(self):
        """Calcula quantas cartas cabem por pÃ¡gina"""
        usable_width = self.a4_width_mm - (2 * self.margin_mm)
        usable_height = self.a4_height_mm - (2 * self.margin_mm)
        
        # Considerar espaÃ§amento entre cartas
        effective_card_width = self.card_width_mm + self.spacing_mm
        effective_card_height = self.card_height_mm + self.spacing_mm
        
        cards_per_row = floor(usable_width / effective_card_width)
        cards_per_column = floor(usable_height / effective_card_height)
        
        return {
            'usable_width': usable_width,
            'usable_height': usable_height,
            'cards_per_row': cards_per_row,
            'cards_per_column': cards_per_column,
            'cards_per_page': cards_per_row * cards_per_column
        }
    
    def get_card_list(self):
        """ObtÃ©m lista de cartas com informaÃ§Ãµes"""
        # Carregar dados do JSON se disponÃ­vel
        cards_data = []
        try:
            with open('cards_database.json', 'r', encoding='utf-8') as f:
                cards_data = json.load(f)
        except:
            pass
        
        # Criar dicionÃ¡rio por imagem
        cards_by_image = {}
        for card in cards_data:
            if 'image' in card:
                image_name = card['image'].split('/')[-1]
                cards_by_image[image_name] = card
        
        # Obter lista de imagens
        cards_folder = './cards'
        image_files = []
        if os.path.exists(cards_folder):
            image_files = [f for f in os.listdir(cards_folder) 
                          if f.lower().endswith('.png')]
            image_files.sort()
        
        return image_files, cards_by_image
    
    def create_document(self, output_file="X-MONSTERS_Cartas_Impressao.docx"):
        """Cria o documento Word com as cartas"""
        layout = self.calculate_layout()
        image_files, cards_by_image = self.get_card_list()
        
        print(f"ğŸ“ Layout calculado:")
        print(f"   - Ãrea Ãºtil: {layout['usable_width']:.1f}mm x {layout['usable_height']:.1f}mm")
        print(f"   - Cartas por linha: {layout['cards_per_row']}")
        print(f"   - Cartas por coluna: {layout['cards_per_column']}")
        print(f"   - Cartas por pÃ¡gina: {layout['cards_per_page']}")
        print(f"ğŸƒ Total de cartas: {len(image_files)}")
        
        total_pages = (len(image_files) + layout['cards_per_page'] - 1) // layout['cards_per_page']
        print(f"ğŸ“„ PÃ¡ginas necessÃ¡rias: {total_pages}")
        
        # Criar documento
        doc = Document()
        
        # Configurar margens
        section = doc.sections[0]
        section.top_margin = Mm(self.margin_mm)
        section.bottom_margin = Mm(self.margin_mm)
        section.left_margin = Mm(self.margin_mm)
        section.right_margin = Mm(self.margin_mm)
        
        current_card = 0
        
        for page in range(total_pages):
            # TÃ­tulo da pÃ¡gina
            if page > 0:
                doc.add_page_break()
            
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run(f'X-MONSTERS - Folha {page + 1}/{total_pages}')
            title_run.font.size = Pt(16)
            title_run.bold = True
            
            # InformaÃ§Ãµes da pÃ¡gina
            cards_on_page = min(layout['cards_per_page'], len(image_files) - current_card)
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run(f'Cartas {current_card + 1} a {current_card + cards_on_page} | DimensÃµes: {self.card_width_mm}mm x {self.card_height_mm}mm')
            info_run.font.size = Pt(10)
            
            # Criar tabela
            table = doc.add_table(rows=layout['cards_per_column'], cols=layout['cards_per_row'])
            table.autofit = False
            
            # Preencher cartas
            for row in range(layout['cards_per_column']):
                table.rows[row].height = Mm(self.card_height_mm)
                
                for col in range(layout['cards_per_row']):
                    if current_card >= len(image_files):
                        break
                    
                    cell = table.cell(row, col)
                    cell.width = Mm(self.card_width_mm)
                    cell._element.clear_content()
                    
                    # Adicionar imagem
                    image_file = image_files[current_card]
                    image_path = os.path.join('./cards', image_file)
                    
                    if os.path.exists(image_path):
                        try:
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            run = paragraph.add_run()
                            run.add_picture(image_path, 
                                          width=Mm(self.card_width_mm), 
                                          height=Mm(self.card_height_mm))
                            
                            # Adicionar nome da carta se disponÃ­vel
                            if image_file in cards_by_image:
                                card_info = cards_by_image[image_file]
                                name_para = cell.add_paragraph()
                                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                name_run = name_para.add_run(card_info.get('name', ''))
                                name_run.font.size = Pt(7)
                                name_run.font.name = 'Arial'
                            
                        except Exception as e:
                            # Erro na imagem
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run(f"âŒ {image_file}")
                            run.font.size = Pt(8)
                            print(f"âš ï¸  Erro ao processar {image_file}: {e}")
                    
                    current_card += 1
                
                if current_card >= len(image_files):
                    break
        
        # Salvar documento
        doc.save(output_file)
        print(f"\nâœ… Documento criado: {output_file}")
        print(f"ğŸ“Š Resumo:")
        print(f"   - Total de cartas processadas: {len(image_files)}")
        print(f"   - Total de pÃ¡ginas: {total_pages}")
        print(f"   - Cartas por pÃ¡gina: {layout['cards_per_page']}")
        print(f"   - DimensÃµes: {self.card_width_mm}mm x {self.card_height_mm}mm")
        
        return output_file

# Exemplo de uso personalizado
if __name__ == "__main__":
    generator = CardPrintingGenerator()
    
    # VocÃª pode modificar estas configuraÃ§Ãµes:
    # generator.card_width_mm = 60     # Largura da carta
    # generator.card_height_mm = 85    # Altura da carta  
    # generator.margin_mm = 15         # Margem das pÃ¡ginas
    # generator.spacing_mm = 3         # EspaÃ§o entre cartas
    
    generator.create_document()