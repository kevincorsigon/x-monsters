from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os
import json
import io

def optimize_image_for_word(image_path, target_width_mm=63, target_height_mm=88, dpi=300):
    """
    Otimiza imagem para melhor qualidade no Word
    """
    # Converter mm para pixels com DPI alto
    width_px = int((target_width_mm / 25.4) * dpi)
    height_px = int((target_height_mm / 25.4) * dpi)
    
    # Abrir e processar imagem
    with Image.open(image_path) as img:
        # Converter para RGB se necessário (evita problemas com CMYK)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Redimensionar com alta qualidade
        img_resized = img.resize((width_px, height_px), Image.Resampling.LANCZOS)
        
        # Salvar como PNG temporário (sem compressão)
        temp_path = image_path.replace('.jpg', '_temp.png').replace('.jpeg', '_temp.png')
        img_resized.save(temp_path, 'PNG', optimize=False)
        
        return temp_path

def create_high_quality_documents():
    """Cria documentos com qualidade otimizada"""
    
    # Configurações de layout
    CARD_WIDTH = Mm(63)
    CARD_HEIGHT = Mm(88)
    CARD_SPACING_MM = 8
    
    # Layout 2x2
    cards_per_row = 2
    cards_per_column = 2
    cards_per_page = cards_per_row * cards_per_column
    
    # Carregar database
    cards_database = []
    cards_by_image = {}
    if os.path.exists('cards_database.json'):
        with open('cards_database.json', 'r', encoding='utf-8') as f:
            cards_database = json.load(f)
        
        for card in cards_database:
            if 'image' in card:
                cards_by_image[card['image']] = card
    
    # Obter lista de imagens
    cards_folder = 'cards'
    if not os.path.exists(cards_folder):
        print("❌ Pasta cards/ não encontrada!")
        return
    
    image_files = [f for f in os.listdir(cards_folder) 
                   if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
    image_files.sort()
    
    total_cards = len(image_files)
    total_pages = (total_cards + cards_per_page - 1) // cards_per_page
    
    print(f"🔧 Otimizando qualidade das imagens...")
    print(f"📊 Total: {total_cards} cartas | {total_pages} páginas")
    
    # Verificar verso
    back_image_path = None
    possible_backs = [
        'verso_card_kevao.png',         # Arquivo PNG preferido
        'verso_card_kevao.jpg',         # Arquivo JPG alternativo
        'verso_card_kevao copiar.jpg'
    ]
    
    for back_name in possible_backs:
        if os.path.exists(back_name):
            back_image_path = back_name
            break
    
    if not back_image_path:
        print("⚠️ Imagem do verso não encontrada")
        return
    
    # Otimizar imagem do verso
    print(f"🎴 Otimizando verso: {back_image_path}")
    optimized_back = optimize_image_for_word(back_image_path)
    
    # Criar documento das frentes
    print("🔄 Criando documento FRENTE (alta qualidade)...")
    doc_front = Document()
    
    # Configurar página
    section = doc_front.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    
    current_card = 0
    
    for page_num in range(total_pages):
        if page_num > 0:
            doc_front.add_page_break()
        
        # Info da página
        cards_on_page = min(cards_per_page, total_cards - current_card)
        info = doc_front.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(f'FRENTES - Página {page_num + 1} | Cartas {current_card + 1}-{current_card + cards_on_page} | ALTA QUALIDADE')
        info_run.font.size = Pt(10)
        info_run.font.bold = True
        
        # Tabela
        table = doc_front.add_table(rows=cards_per_column, cols=cards_per_row)
        table.autofit = False
        
        for row in range(cards_per_column):
            table.rows[row].height = CARD_HEIGHT
            
            for col in range(cards_per_row):
                if current_card >= total_cards:
                    break
                
                cell = table.cell(row, col)
                cell.width = CARD_WIDTH
                cell._element.clear_content()
                
                # Otimizar imagem da carta
                image_file = image_files[current_card]
                image_path = os.path.join(cards_folder, image_file)
                
                if os.path.exists(image_path):
                    try:
                        print(f"🖼️ Otimizando {image_file} ({current_card + 1}/{total_cards})")
                        
                        # Otimizar imagem
                        optimized_path = optimize_image_for_word(image_path)
                        
                        paragraph = cell.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = paragraph.add_run()
                        run.add_picture(optimized_path, width=CARD_WIDTH, height=CARD_HEIGHT)
                        
                        # Limpar arquivo temporário
                        if os.path.exists(optimized_path):
                            os.remove(optimized_path)
                        
                        # Nome da carta
                        if image_file in cards_by_image:
                            card_info = cards_by_image[image_file]
                            name_para = cell.add_paragraph()
                            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            name_run = name_para.add_run(card_info.get('name', ''))
                            name_run.font.size = Pt(7)
                        
                    except Exception as e:
                        print(f"❌ Erro ao otimizar {image_file}: {e}")
                        # Usar imagem original como fallback
                        try:
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=CARD_WIDTH, height=CARD_HEIGHT)
                        except:
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run(f"❌ {image_file}")
                            run.font.size = Pt(8)
                
                current_card += 1
    
    # Salvar frentes
    front_filename = "X-MONSTERS_Frentes_ALTA_QUALIDADE.docx"
    doc_front.save(front_filename)
    print(f"✅ Frentes salvas: {front_filename}")
    
    # Criar documento dos versos
    print("🔄 Criando documento VERSO (alta qualidade)...")
    doc_back = Document()
    
    # Configurar página
    section = doc_back.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    
    for page_num in range(total_pages):
        if page_num > 0:
            doc_back.add_page_break()
        
        # Info da página
        info = doc_back.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(f'VERSOS - Página {page_num + 1} | ALTA QUALIDADE')
        info_run.font.size = Pt(10)
        info_run.font.bold = True
        
        # Tabela
        table = doc_back.add_table(rows=cards_per_column, cols=cards_per_row)
        table.autofit = False
        
        for row in range(cards_per_column):
            table.rows[row].height = CARD_HEIGHT
            
            for col in range(cards_per_row):
                cell = table.cell(row, col)
                cell.width = CARD_WIDTH
                cell._element.clear_content()
                
                try:
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = paragraph.add_run()
                    run.add_picture(optimized_back, width=CARD_WIDTH, height=CARD_HEIGHT)
                    
                except Exception as e:
                    print(f"❌ Erro no verso: {e}")
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run("❌ Verso")
                    run.font.size = Pt(8)
    
    # Salvar versos
    back_filename = "X-MONSTERS_Versos_ALTA_QUALIDADE.docx"
    doc_back.save(back_filename)
    print(f"✅ Versos salvos: {back_filename}")
    
    # Limpar arquivo temporário do verso
    if os.path.exists(optimized_back):
        os.remove(optimized_back)
    
    print("\n🎉 DOCUMENTOS ALTA QUALIDADE CRIADOS:")
    print(f"📄 Frentes: {front_filename}")
    print(f"📄 Versos: {back_filename}")
    print("🔧 Otimizações aplicadas:")
    print("   • Conversão para PNG (sem compressão)")
    print("   • Redimensionamento com algoritmo LANCZOS")
    print("   • DPI alto (300)")
    print("   • Conversão RGB para compatibilidade")

if __name__ == "__main__":
    create_high_quality_documents()