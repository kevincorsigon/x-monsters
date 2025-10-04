from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
from math import floor

def create_cards_with_backs():
    # DimensÃµes das cartas
    CARD_WIDTH = Mm(63)
    CARD_HEIGHT = Mm(88)
    
    # ConfiguraÃ§Ãµes A4
    A4_WIDTH_MM = 210
    A4_HEIGHT_MM = 297
    MARGIN_MM = 15  # Margem maior
    CARD_SPACING_MM = 8  # EspaÃ§o entre cartas para facilitar recorte
    
    # Calcular layout com espaÃ§amento maior
    usable_width_mm = A4_WIDTH_MM - (2 * MARGIN_MM)
    usable_height_mm = A4_HEIGHT_MM - (2 * MARGIN_MM)
    
    # Considerar espaÃ§amento entre cartas
    effective_card_width = 63 + CARD_SPACING_MM
    effective_card_height = 88 + CARD_SPACING_MM
    
    cards_per_row = floor(usable_width_mm / effective_card_width)
    cards_per_column = floor(usable_height_mm / effective_card_height)
    cards_per_page = cards_per_row * cards_per_column
    
    print(f"ğŸ“ Layout com margem maior:")
    print(f"   - Ãrea Ãºtil A4: {usable_width_mm:.1f}mm x {usable_height_mm:.1f}mm")
    print(f"   - EspaÃ§amento entre cartas: {CARD_SPACING_MM}mm")
    print(f"   - Cartas por linha: {cards_per_row}")
    print(f"   - Cartas por coluna: {cards_per_column}")
    print(f"   - Cartas por pÃ¡gina: {cards_per_page}")
    
    # Carregar dados das cartas
    try:
        with open('cards_database.json', 'r', encoding='utf-8') as f:
            cards_data = json.load(f)
        print(f"âœ… Carregadas {len(cards_data)} cartas do database")
    except:
        print("âš ï¸  Erro ao carregar cards_database.json")
        cards_data = []
    
    # Criar dicionÃ¡rio por imagem
    cards_by_image = {}
    for card in cards_data:
        if 'image' in card:
            image_name = card['image'].split('/')[-1]
            cards_by_image[image_name] = card
    
    # Obter lista de imagens
    cards_folder = './cards'
    if os.path.exists(cards_folder):
        image_files = [f for f in os.listdir(cards_folder) if f.lower().endswith('.png')]
        image_files.sort()
        print(f"ğŸ“ Encontradas {len(image_files)} imagens na pasta cards/")
    else:
        print("âŒ Pasta cards/ nÃ£o encontrada!")
        return
    
    # Verificar se existe imagem do verso
    back_image_path = None
    possible_backs = [
        'verso_card_kevao.png',         # Arquivo PNG preferido
        'verso_card_kevao.jpg',         # Arquivo JPG alternativo
        'verso_card_kevao copiar.jpg',  # Nome alternativo
        'verso_card.png', 
        'card_back.png'
    ]
    
    for back_name in possible_backs:
        if os.path.exists(back_name):
            back_image_path = back_name
            break
        elif os.path.exists(f'cards/{back_name}'):
            back_image_path = f'cards/{back_name}'
            break
    
    if back_image_path:
        print(f"ğŸ´ Verso encontrado: {back_image_path}")
    else:
        print("âš ï¸  Verso nÃ£o encontrado. Procurando por: verso_card_kevao.png")
        # Criar verso simples se nÃ£o encontrar
        create_simple_back = True
    
    total_cards = len(image_files)
    total_pages = (total_cards + cards_per_page - 1) // cards_per_page
    
    print(f"ğŸ“Š Gerando documentos:")
    print(f"   - Total de cartas: {total_cards}")
    print(f"   - PÃ¡ginas por documento: {total_pages}")
    
    # ===== DOCUMENTO DA FRENTE =====
    print("\nğŸ”„ Criando documento da FRENTE das cartas...")
    doc_front = Document()
    
    # Configurar margens
    section = doc_front.sections[0]
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    
    current_card = 0
    
    for page in range(total_pages):
        if page > 0:
            doc_front.add_page_break()
        
        # TÃ­tulo
        title = doc_front.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f'X-MONSTERS - FRENTES - PÃ¡gina {page + 1}/{total_pages}')
        title_run.font.size = Pt(16)
        title_run.bold = True
        
        # Info
        cards_on_page = min(cards_per_page, total_cards - current_card)
        info = doc_front.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(f'Cartas {current_card + 1} a {current_card + cards_on_page} | 63mmÃ—88mm | Margem: {CARD_SPACING_MM}mm')
        info_run.font.size = Pt(10)
        
        # Criar tabela
        table = doc_front.add_table(rows=cards_per_column, cols=cards_per_row)
        table.autofit = False
        
        for row in range(cards_per_column):
            table.rows[row].height = CARD_HEIGHT
            
            for col in range(cards_per_row):
                if current_card >= total_cards:
                    break
                
                cell = table.cell(row, col)
                cell.width = CARD_WIDTH
                cell._element.clear_content()
                
                # Adicionar imagem da carta
                image_file = image_files[current_card]
                image_path = os.path.join(cards_folder, image_file)
                
                if os.path.exists(image_path):
                    try:
                        paragraph = cell.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=CARD_WIDTH, height=CARD_HEIGHT)
                        
                        # Nome da carta
                        if image_file in cards_by_image:
                            card_info = cards_by_image[image_file]
                            name_para = cell.add_paragraph()
                            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            name_run = name_para.add_run(card_info.get('name', ''))
                            name_run.font.size = Pt(7)
                        
                    except Exception as e:
                        paragraph = cell.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(f"âŒ {image_file}")
                        run.font.size = Pt(8)
                        print(f"âš ï¸  Erro ao processar {image_file}: {e}")
                
                current_card += 1
            
            if current_card >= total_cards:
                break
    
    # Salvar documento da frente
    front_file = 'X-MONSTERS_Frentes_Impressao.docx'
    doc_front.save(front_file)
    print(f"âœ… Frentes salvas: {front_file}")
    
    # ===== DOCUMENTO DO VERSO =====
    print("\nğŸ”„ Criando documento do VERSO das cartas...")
    doc_back = Document()
    
    # Configurar margens
    section = doc_back.sections[0]
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    
    for page in range(total_pages):
        if page > 0:
            doc_back.add_page_break()
        
        # TÃ­tulo
        title = doc_back.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f'X-MONSTERS - VERSOS - PÃ¡gina {page + 1}/{total_pages}')
        title_run.font.size = Pt(16)
        title_run.bold = True
        
        # Info importante para impressÃ£o
        info = doc_back.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run('âš ï¸ IMPRIMIR NO VERSO DO PAPEL DA PÃGINA CORRESPONDENTE âš ï¸')
        info_run.font.size = Pt(12)
        info_run.bold = True
        
        # Criar tabela
        table = doc_back.add_table(rows=cards_per_column, cols=cards_per_row)
        table.autofit = False
        
        cards_on_this_page = min(cards_per_page, total_cards - (page * cards_per_page))
        
        for row in range(cards_per_column):
            table.rows[row].height = CARD_HEIGHT
            
            for col in range(cards_per_row):
                card_index = (page * cards_per_page) + (row * cards_per_row) + col
                
                if card_index >= total_cards:
                    break
                
                cell = table.cell(row, col)
                cell.width = CARD_WIDTH
                cell._element.clear_content()
                
                # Adicionar verso (mesmo verso para todas as cartas)
                try:
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if back_image_path and os.path.exists(back_image_path):
                        run = paragraph.add_run()
                        run.add_picture(back_image_path, width=CARD_WIDTH, height=CARD_HEIGHT)
                    else:
                        # Criar verso simples de texto se nÃ£o tiver imagem
                        run = paragraph.add_run("X-MONSTERS")
                        run.font.size = Pt(20)
                        run.font.bold = True
                        
                        # Adicionar decoraÃ§Ã£o
                        deco_para = cell.add_paragraph()
                        deco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        deco_run = deco_para.add_run("ğŸ® CARD GAME ğŸ®")
                        deco_run.font.size = Pt(12)
                
                except Exception as e:
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run("VERSO")
                    run.font.size = Pt(16)
                    print(f"âš ï¸  Erro ao processar verso: {e}")
            
            if (page * cards_per_page) + (row * cards_per_row) >= total_cards:
                break
    
    # Salvar documento do verso
    back_file = 'X-MONSTERS_Versos_Impressao.docx'
    doc_back.save(back_file)
    print(f"âœ… Versos salvos: {back_file}")
    
    print(f"\nğŸ‰ DOCUMENTOS CRIADOS:")
    print(f"   ğŸ“„ Frentes: {front_file}")
    print(f"   ğŸ“„ Versos: {back_file}")
    print(f"   ğŸ¯ Layout: {cards_per_row}Ã—{cards_per_column} por pÃ¡gina")
    print(f"   ğŸ“ Margem entre cartas: {CARD_SPACING_MM}mm")
    print(f"   ğŸ“ Total de pÃ¡ginas: {total_pages}")
    
    return front_file, back_file

if __name__ == "__main__":
    create_cards_with_backs()